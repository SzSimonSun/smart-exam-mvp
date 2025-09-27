"""
智能文档处理服务 - 实用版本
处理PDF、DOCX、图片文件，将试卷拆分为单个题目
"""

import io
import json
import logging
import uuid
import re
import tempfile
from typing import List, Dict, Any
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """智能文档处理器 - 支持AI增强识别"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'image/jpeg': self._process_image,
            'image/jpg': self._process_image,
            'image/png': self._process_image,
        }
        
        # AI增强配置
        self.use_ai_enhancement = False  # 可配置是否启用AI增强
        self.ai_api_key = None  # AI API密钥
        self.min_confidence_threshold = 0.6  # 最低置信度阈值
        self.min_questions_threshold = 3  # 最少题目数量阈值
        
        # 检查依赖
        self.has_pdf = self._check_pdf_support()
        self.has_docx = self._check_docx_support()
        
        # 尝试初始化AI功能
        self._try_initialize_ai()
        
        # AI增强配置
        self.use_ai_enhancement = False  # 可配置是否启用AI增强
        self.ai_api_key = None  # AI API密钥
        self.min_confidence_threshold = 0.6  # 最低置信度阈值
        self.min_questions_threshold = 3  # 最少题目数量阈值
        
        # 检查依赖
        self.has_pdf = self._check_pdf_support()
        self.has_docx = self._check_docx_support()
        
        # 尝试初始化AI功能
        self._try_initialize_ai()
        
        logger.info(f"文档处理器初始化 - PDF支持: {self.has_pdf}, DOCX支持: {self.has_docx}, AI增强: {self.use_ai_enhancement}")
    
    def _try_initialize_ai(self):
        """尝试初始化AI功能"""
        try:
            # 检查环境变量中是否有API密钥
            import os
            api_key = os.getenv('OPENAI_API_KEY') or os.getenv('AI_API_KEY')
            
            if api_key:
                self.ai_api_key = api_key
                self.use_ai_enhancement = True
                logger.info("AI增强功能已启用")
            else:
                logger.info("AI API密钥未配置，使用传统方法")
                
        except Exception as e:
            logger.warning(f"AI初始化失败: {e}")
            self.use_ai_enhancement = False
    
    def _should_use_ai_enhancement(self, regex_questions: List[Dict[str, Any]]) -> bool:
        """判断是否需要使用AI增强"""
        if not self.use_ai_enhancement:
            return False
            
        # 检查数量和质量
        if len(regex_questions) < self.min_questions_threshold:
            logger.info(f"题目数量过少 ({len(regex_questions)} < {self.min_questions_threshold})，启用AI增强")
            return True
            
        # 检查平均置信度
        if regex_questions:
            avg_confidence = sum(q.get('confidence', 0) for q in regex_questions) / len(regex_questions)
            if avg_confidence < self.min_confidence_threshold:
                logger.info(f"平均置信度过低 ({avg_confidence:.2f} < {self.min_confidence_threshold})，启用AI增强")
                return True
        
        return False
    
    def _ai_enhanced_split(self, text: str) -> List[Dict[str, Any]]:
        """使用AI增强进行文档拆分"""
        try:
            # 这里可以集成不同的AI服务
            if self._is_openai_available():
                return self._call_openai_api(text)
            else:
                # 退而求其次，使用增强正则处理器
                logger.warning("AI服务不可用，使用增强正则处理器")
                return self._enhanced_regex_split(text)
                
        except Exception as e:
            logger.error(f"AI增强处理失败: {e}")
            return self._enhanced_regex_split(text)
    
    def _is_openai_available(self) -> bool:
        """检查OpenAI是否可用"""
        try:
            import openai
            return self.ai_api_key is not None
        except ImportError:
            return False
    
    def _call_openai_api(self, text: str) -> List[Dict[str, Any]]:
        """调用OpenAI API进行文档分析"""
        try:
            import openai
            
            client = openai.OpenAI(api_key=self.ai_api_key)
            
            prompt = f"""
请分析以下考试试卷文档，准确识别和拆分每道题目。

要求：
1. 识别每道题目的完整内容（包括题号、题干、选项、填空等）
2. 判断题目类型：single(单选)、multiple(多选)、fill(填空)、judge(判断)、subjective(主观)
3. 保持原有格式和排版
4. 给出识别置信度(0-1之间的小数)
5. 建议相关知识点

输出格式：
```json
[
    {{
        "seq": 1,
        "question_text": "完整的题目内容，保持原格式",
        "question_type": "single",
        "confidence": 0.95,
        "suggested_kps": ["相关知识点1", "相关知识点2"]
    }}
]
```

文档内容：
{text[:3000]}  # 限制长度避免超出token限制
"""
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",  # 使用性价比更高的模型
                messages=[
                    {"role": "system", "content": "你是一个专业的试卷分析专家，擅长准确识别和拆分考试题目。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=4000
            )
            
            result_text = response.choices[0].message.content
            return self._parse_ai_response(result_text)
            
        except Exception as e:
            logger.error(f"OpenAI API调用失败: {e}")
            return []
    
    def _parse_ai_response(self, response_text: str) -> List[Dict[str, Any]]:
        """解析AI响应"""
        try:
            import json
            import re
            
            # 提取JSON部分
            json_match = re.search(r'```json\n(.*?)\n```', response_text, re.DOTALL)
            if json_match:
                json_text = json_match.group(1)
            else:
                # 如果没有代码块，尝试直接解析
                json_text = response_text.strip()
            
            questions = json.loads(json_text)
            
            # 验证和标准化结果
            validated_questions = []
            for i, q in enumerate(questions):
                if self._validate_ai_question(q):
                    # 为了兼容现有系统，添加必要的字段
                    validated_q = {
                        'seq': q.get('seq', i + 1),
                        'question_text': q.get('question_text', ''),
                        'question_type': q.get('question_type', 'single'),
                        'confidence': q.get('confidence', 0.8),
                        'ocr_result': {
                            'text': q.get('question_text', ''),
                            'type': q.get('question_type', 'single'),
                            'confidence': q.get('confidence', 0.8)
                        },
                        'candidate_kps': q.get('suggested_kps', []),
                        'review_status': 'pending'
                    }
                    validated_questions.append(validated_q)
            
            logger.info(f"AI解析成功，识别 {len(validated_questions)} 道题目")
            return validated_questions
            
        except Exception as e:
            logger.error(f"解析AI响应失败: {e}")
            return []
    
    def _validate_ai_question(self, question: Dict) -> bool:
        """验证AI返回的题目"""
        required_fields = ['question_text', 'question_type']
        return all(field in question for field in required_fields) and \
               len(question.get('question_text', '').strip()) > 10
    
    def _enhanced_regex_split(self, text: str) -> List[Dict[str, Any]]:
        """增强正则表达式拆分（用作退备方案）"""
        # 这里可以集成之前实现的增强正则处理器
        # 暂时使用现有的方法
        return self._regex_split_text(text)
    
    def _check_pdf_support(self) -> bool:
        try:
            import PyPDF2
            return True
        except ImportError:
            return False
    
    def _check_docx_support(self) -> bool:
        try:
            from docx import Document
            return True
        except ImportError:
            return False
    
    def process_document(self, file_content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
        """处理文档，拆分为题目"""
        if content_type not in self.supported_types:
            return {
                'success': False,
                'error': f"不支持的文件类型: {content_type}",
                'items': []
            }
        
        try:
            processor = self.supported_types[content_type]
            result = processor(file_content, filename)
            
            # 为每个题目添加元数据
            for i, item in enumerate(result['items']):
                item['seq'] = i + 1
                item['id'] = str(uuid.uuid4())
                item['source_file'] = filename
                
            return result
            
        except Exception as e:
            logger.error(f"处理文档失败 {filename}: {e}")
            return {
                'success': False,
                'error': str(e),
                'items': []
            }
    
    def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """处理PDF文件"""
        if not self.has_pdf:
            return self._process_pdf_fallback(filename)
        
        try:
            import PyPDF2
            logger.info(f"开始处理PDF文件: {filename}, 大小: {len(file_content)} bytes")
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            full_text = ""
            page_count = len(pdf_reader.pages)
            
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    # 保持页面的结构和格式
                    if page_text.strip():
                        # 在页面之间添加适当的分隔
                        if i > 0:
                            full_text += '\n\n'
                        full_text += self._preserve_pdf_format(page_text)
                except Exception as e:
                    logger.warning(f"第{i+1}页文本提取失败: {e}")
                    continue
            
            if not full_text.strip():
                return {
                    'success': False,
                    'error': 'PDF文本提取为空，可能是图片扫描版',
                    'items': []
                }
            
            questions = self._smart_split_text(full_text)
            
            return {
                'success': True,
                'file_type': 'pdf',
                'total_pages': page_count,
                'items': questions,
                'message': f'从PDF中提取了 {len(questions)} 道题目'
            }
            
        except Exception as e:
            logger.error(f"PDF处理失败: {e}")
            return self._process_pdf_fallback(filename)
    
    def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """处理DOCX文件"""
        if not self.has_docx:
            return self._process_docx_fallback(filename)
        
        try:
            from docx import Document
            logger.info(f"开始处理DOCX文件: {filename}")
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                
                doc = Document(tmp_file.name)
                full_text = ""
                
                # 保持段落结构和格式
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        # 保持段落的缩进和对齐
                        para_text = paragraph.text
                        
                        # 检查是否为标题或项目符号
                        if self._is_title_or_numbering(para_text):
                            full_text += '\n' + para_text + '\n'
                        else:
                            full_text += para_text + '\n'
                
                # 处理表格内容（如果有）
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        full_text += '\n' + table_text + '\n'
            
            Path(tmp_file.name).unlink(missing_ok=True)
            
            if not full_text.strip():
                return {
                    'success': False,
                    'error': '文档内容为空',
                    'items': []
                }
            
            questions = self._smart_split_text(full_text)
            
            return {
                'success': True,
                'file_type': 'docx',
                'items': questions,
                'message': f'从DOCX中提取了 {len(questions)} 道题目'
            }
            
        except Exception as e:
            logger.error(f"DOCX处理失败: {e}")
            return self._process_docx_fallback(filename)
    
    def _process_image(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """处理图片文件（使用模拟OCR）"""
        try:
            logger.info(f"处理图片文件: {filename}")
            
            # 模拟OCR识别结果
            mock_text = f"""
1. 下列哪个数是偶数？
A. 1  B. 3  C. 6  D. 7

2. 计算：3 × 4 + 2 = ______

3. 判断：正方形的四条边都相等。（  ）

（注：这是从图片 {filename} 模拟识别的内容）
            """.strip()
            
            questions = self._smart_split_text(mock_text)
            
            return {
                'success': True,
                'file_type': 'image',
                'items': questions,
                'message': f'从图片中识别了 {len(questions)} 道题目（模拟）'
            }
            
        except Exception as e:
            logger.error(f"图片处理失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'items': []
            }
    
    def _smart_split_text(self, text: str) -> List[Dict[str, Any]]:
        """智能文本拆分 - 改进版（保持格式 + AI增强）"""
        logger.info(f"开始智能拆分文本，长度: {len(text)} 字符")
        
        # 首先清理文本，但保留重要的格式信息
        cleaned_text = self._clean_text_preserve_format(text)
        logger.info(f"清理后文本长度: {len(cleaned_text)} 字符")
        
        # 传统正则表达式处理
        regex_questions = self._regex_split_text(cleaned_text)
        
        # 评估传统方法的效果
        should_use_ai = self._should_use_ai_enhancement(regex_questions)
        
        if should_use_ai and self.use_ai_enhancement:
            logger.info("传统方法效果不佳，尝试AI增强")
            ai_questions = self._ai_enhanced_split(cleaned_text)
            
            # 选择最佳结果
            if len(ai_questions) > len(regex_questions):
                logger.info(f"AI增强效果更好：{len(ai_questions)} vs {len(regex_questions)}")
                return ai_questions
        
        return regex_questions
    
    def _regex_split_text(self, cleaned_text: str) -> List[Dict[str, Any]]:
        """传统正则表达式拆分"""
        questions = []
        
        # 改进的题号匹配模式，保持原始格式
        patterns = [
            # 最精确的模式，保留空格和换行
            r'(?:^|\n)\s*([1-9]\d*)[.\uff0e]\s*([^\n]+(?:\n(?!\s*\d+[.\uff0e]|\s*[一二三四五六七八九十]、)[^\n]*)*?)(?=\n\s*\d+[.\uff0e]|\n\s*[一二三四五六七八九十]、|$)',
            r'(?:^|\n)\s*([1-9]\d*)、\s*([^\n]+(?:\n(?!\s*\d+、|\s*[一二三四五六七八九十]、)[^\n]*)*?)(?=\n\s*\d+、|\n\s*[一二三四五六七八九十]、|$)',
            r'(?:^|\n)\s*第\s*([1-9]\d*)\s*题[\uff1a:\s]*([^\n]+(?:\n(?!\s*第\s*\d+\s*题|\s*[一二三四五六七八九十]、)[^\n]*)*?)(?=\n\s*第\s*\d+\s*题|\n\s*[一二三四五六七八九十]、|$)',
            r'(?:^|\n)\s*([1-9]\d*)\)\s*([^\n]+(?:\n(?!\s*\d+\)|\s*[一二三四五六七八九十]、)[^\n]*)*?)(?=\n\s*\d+\)|\n\s*[一二三四五六七八九十]、|$)',
            r'(?:^|\n)\s*\(([1-9]\d*)\)\s*([^\n]+(?:\n(?!\s*\(\d+\)|\s*[一二三四五六七八九十]、)[^\n]*)*?)(?=\n\s*\(\d+\)|\n\s*[一二三四五六七八九十]、|$)',
        ]
        
        best_questions = []
        best_pattern = None
        
        for pattern in patterns:
            matches = re.findall(pattern, cleaned_text, re.MULTILINE | re.DOTALL)
            if matches:
                temp_questions = []
                for num, content in matches:
                    # 保持原始格式的内容处理
                    formatted_content = self._preserve_question_format(content.strip())
                    # 更严格的内容过滤
                    if self._is_valid_question(formatted_content):
                        temp_questions.append(self._create_question_item(formatted_content, int(num)))
                
                if len(temp_questions) > len(best_questions):
                    best_questions = temp_questions
                    best_pattern = pattern
        
        if best_questions:
            logger.info(f"使用模式匹配找到 {len(best_questions)} 道题目")
            return best_questions
        
        # 如果没找到带题号的，尝试其他策略
        return self._fallback_split_strategies(cleaned_text)
    
    def _clean_text_preserve_format(self, text: str) -> str:
        """清理文本，保留重要的格式信息"""
        # 移除干扰内容，但保留题目内部的格式
        noise_patterns = [
            r'请将信息填写清楚.*?不得有误',
            r'保持答题卡.*?不得折叠',
            r'选择题用.*?橡皮擦',
            r'姓名.*?准考证号',
            r'考场号.*?座位号',
            r'注意事项.*?(?=\d+[.\uff0e]|$)',
            r'答题说明.*?(?=\d+[.\uff0e]|$)',
            r'考试时间.*?分钟',
            r'满分.*?分',
            r'《.*?》',  # 移除标题
            r'^\s*[一二三四五六七八九十][、.].*?题.*?分.*?$',  # 移除大题头
        ]
        
        cleaned_text = text
        for pattern in noise_patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.MULTILINE | re.IGNORECASE)
        
        # 保留适度的空行，但避免过多
        cleaned_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned_text)
        
        # 保持重要的段落结构
        cleaned_text = self._preserve_paragraph_structure(cleaned_text)
        
        return cleaned_text.strip()
    
    def _preserve_paragraph_structure(self, text: str) -> str:
        """保持段落结构"""
        lines = text.split('\n')
        structured_lines = []
        
        for i, line in enumerate(lines):
            line = line.rstrip()
            
            # 保持题号前的空行
            if re.match(r'^\s*\d+[.\uff0e]', line) and i > 0:
                if structured_lines and structured_lines[-1].strip():
                    structured_lines.append('')  # 添加空行
            
            structured_lines.append(line)
        
        return '\n'.join(structured_lines)
    
    def _preserve_question_format(self, content: str) -> str:
        """保持题目内部的格式"""
        # 保留选择项的对齐格式
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.rstrip()  # 移除行尾空格
            
            # 对选择项进行格式化，增加缩进
            if re.match(r'^\s*[A-Z][.\uff0e、]', line):
                # 保持选择项的缩进，使用四个空格
                clean_line = line.strip()
                formatted_lines.append('    ' + clean_line)
            # 对数学公式和表达式进行特殊处理
            elif re.search(r'[=+\-×÷\(\)\{\}\[\]]', line):
                # 保持数学公式的缩进（如果本身有缩进）
                if line.startswith('   '):
                    formatted_lines.append('    ' + line.strip())
                else:
                    formatted_lines.append(line)
            else:
                formatted_lines.append(line)
        
        # 重新组合并保持换行
        formatted_content = '\n'.join(formatted_lines)
        
        # 保持常见的数学公式格式
        formatted_content = self._preserve_math_format(formatted_content)
        
        # 保持填空题的下划线格式
        formatted_content = self._preserve_fill_blank_format(formatted_content)
        
        return formatted_content
    
    def _preserve_fill_blank_format(self, text: str) -> str:
        """保持填空题的下划线格式"""
        # 保持下划线的格式
        text = re.sub(r'_{3,}', '______', text)  # 标准化下划线长度
        
        # 保持省略号的格式
        text = re.sub(r'…+', '……', text)  # 标准化省略号
        
        # 保持括号的格式
        text = re.sub(r'\(　*\)', '（　　）', text)  # 标准化括号格式
        
        return text
    
    def _preserve_math_format(self, text: str) -> str:
        """保持数学公式的格式"""
        # 保持常见的数学符号和格式
        math_replacements = [
            (r'\s*\^\s*', '^'),  # 上标
            (r'\s*_\s*', '_'),   # 下标
            (r'\s*=\s*', ' = '), # 等号周围的空格
            (r'\s*\+\s*', ' + '), # 加号
            (r'\s*-\s*', ' - '), # 减号
            (r'\s*\*\s*', ' × '), # 乘号
            (r'\s*/\s*', ' ÷ '), # 除号
        ]
        
        for pattern, replacement in math_replacements:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _is_title_or_numbering(self, text: str) -> bool:
        """判断是否为标题或编号"""
        # 检查是否为章节标题或项目编号
        title_patterns = [
            r'^第.*章',  # 第一章
            r'^第.*节',  # 第一节
            r'^[\u4e00-\u4e5d][、.]',  # 一、
            r'^\d+[\.、]\s*[^\d]',  # 1.
            r'^选择题|填空题|判断题|主观题',  # 题型标题
        ]
        
        for pattern in title_patterns:
            if re.match(pattern, text.strip()):
                return True
        return False
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本并保持结构"""
        table_text = ""
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text += '  '.join(row_text) + '\n'
        return table_text
    
    def _preserve_pdf_format(self, text: str) -> str:
        """保持PDF文本的格式"""
        # PDF提取的文本可能格式混乱，需要整理
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 保持题号格式
            if re.match(r'^\s*\d+[.\uff0e]', line):
                formatted_lines.append('\n' + line)  # 题号前加空行
            # 保持选择项格式
            elif re.match(r'^\s*[A-Z][.\uff0e、]', line):
                formatted_lines.append('  ' + line)  # 选择项缩进
            else:
                formatted_lines.append(line)
        
        return '\n'.join(formatted_lines)
        """清理文本，移除干扰内容"""
        # 移除常见的答题卡说明
        noise_patterns = [
            r'请将信息填写清楚.*?不得有误',
            r'保持答题卡.*?不得折叠',
            r'选择题用.*?橡皮擦',
            r'姓名.*?准考证号',
            r'考场号.*?座位号',
            r'注意事项.*?(?=\d+[.\uff0e]|$)',
            r'答题说明.*?(?=\d+[.\uff0e]|$)',
            r'考试时间.*?分钟',
            r'满分.*?分',
            r'^\s*[A-D]\[\]\s*[A-D]\[\].*$',  # 纯选项框
            r'^\s*_{5,}\s*$',  # 纯下划线
            r'^\s*\([^)]*分\)\s*$',  # 纯分数标记
        ]
        
        cleaned_text = text
        for pattern in noise_patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.MULTILINE | re.IGNORECASE)
        
        # 移除多余的空行
        cleaned_text = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_text)
        
        return cleaned_text.strip()
    
    def _is_valid_question(self, content: str) -> bool:
        """判断是否是有效的题目内容"""
        content = content.strip()
        
        # 长度检查
        if len(content) < 15:  # 提高最小长度要求
            return False
        
        # 排除纯格式内容
        invalid_patterns = [
            r'^\s*[A-D]\[\]\s*[A-D]\[\]',  # 纯选项框
            r'^\s*_{5,}',  # 以下划线开头
            r'^\s*\d+\s*[.\uff0e]\s*[A-D]\[\]',  # 题号+选项框
            r'^\s*\([^)]*分\)',  # 纯分数
            r'^\s*第\s*\d+\s*页',  # 页码
            r'^\s*共\s*\d+\s*页',  # 页码信息
        ]
        
        for pattern in invalid_patterns:
            if re.match(pattern, content):
                return False
        
        # 必须包含实际问题内容的指标
        question_indicators = [
            r'[?？]',  # 问号
            r'下列.*?是|哪.*?是|什么.*?是',  # 问题词
            r'计算|求|证明|分析|简述|论述',  # 题目动词
            r'正确.*?是|错误.*?是',  # 判断题特征
            r'填入.*?空|填空|完成.*?句子',  # 填空题特征
        ]
        
        has_question_feature = any(re.search(pattern, content, re.IGNORECASE) for pattern in question_indicators)
        
        # 或者包含完整的选择题格式
        has_choices = bool(re.search(r'[A-D][.\uff0e].*?[A-D][.\uff0e]', content))
        
        return has_question_feature or has_choices
    
    def _fallback_split_strategies(self, text: str) -> List[Dict[str, Any]]:
        """备用拆分策略"""
        logger.info("使用备用拆分策略")
        
        questions = []
        
        # 策略1：基于问号分割
        question_parts = re.split(r'[?？]\s*(?=\n|$)', text)
        for i, part in enumerate(question_parts[:-1]):  # 最后一部分通常是不完整的
            part = part.strip()
            if self._is_valid_question(part + '？'):
                questions.append(self._create_question_item(part + '？', i + 1))
        
        if questions:
            logger.info(f"问号分割策略找到 {len(questions)} 道题目")
            return questions
        
        # 策略2：基于选择题特征分割
        choice_pattern = r'([^\n]*[A-D][.\uff0e][^\n]*(?:\n[^\n]*[A-D][.\uff0e][^\n]*)*(?:\n[^\n]*[A-D][.\uff0e][^\n]*)*(?:\n[^\n]*[A-D][.\uff0e][^\n]*)*)'
        choice_matches = re.findall(choice_pattern, text, re.MULTILINE)
        
        for i, match in enumerate(choice_matches):
            if self._is_valid_question(match):
                questions.append(self._create_question_item(match.strip(), i + 1))
        
        if questions:
            logger.info(f"选择题特征分割找到 {len(questions)} 道题目")
            return questions
        
        # 策略3：如果前面都失败，创建一个包含所有内容的题目
        if text.strip() and len(text.strip()) > 50:
            questions.append(self._create_question_item(text.strip(), 1))
            logger.info("创建单一综合题目")
        
        return questions
    
    def _create_question_item(self, question_text: str, seq: int) -> Dict[str, Any]:
        """创建题目项"""
        question_type = self._detect_question_type(question_text)
        confidence = self._calculate_confidence(question_text, question_type)
        
        return {
            'question_text': question_text.strip(),
            'question_type': question_type,
            'confidence': confidence,
            'ocr_result': {
                'text': question_text.strip(),
                'type': question_type,
                'confidence': confidence
            },
            'candidate_kps': self._suggest_knowledge_points(question_text),
            'review_status': 'pending'
        }
    
    def _detect_question_type(self, text: str) -> str:
        """检测题目类型 - 改进版"""
        text_lower = text.lower()
        
        # 判断题 - 更精确的匹配
        judge_patterns = [
            r'判断.*?题',
            r'对错.*?题',
            r'正确.*?是|错误.*?是',
            r'是否.*?正确',
            r'\(　*\)\s*$',  # 结尾的括号
            r'说法.*?正确',
        ]
        
        for pattern in judge_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return 'judge'
        
        # 填空题 - 更精确的匹配
        fill_patterns = [
            r'_{3,}',  # 三个或以上下划线
            r'填空.*?题',
            r'完成.*?句子',
            r'填入.*?空白',
            r'……',  # 省略号
            r'　+',  # 空格
        ]
        
        for pattern in fill_patterns:
            if re.search(pattern, text):
                return 'fill'
        
        # 选择题 - 区分单选和多选
        choice_pattern = r'[A-D][.\uff0e、]'
        choice_count = len(re.findall(choice_pattern, text))
        
        if choice_count >= 2:
            # 检查是否是多选题
            multi_patterns = [
                r'多选.*?题',
                r'多项.*?选择',
                r'哪些.*?正确',
                r'以下.*?哪些',
                r'包括.*?哪些',
            ]
            
            for pattern in multi_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    return 'multiple'
            
            return 'single'
        
        # 主观题 - 更全面的检测
        subjective_patterns = [
            r'论述.*?题',
            r'分析.*?题',
            r'简答.*?题',
            r'计算.*?题',
            r'证明.*?题',
            r'请.*?论述',
            r'请.*?分析',
            r'请.*?说明',
            r'为什么',
            r'怎么样',
            r'如何.*?理解',
            r'谈谈.*?看法',
            r'结合.*?说明',
        ]
        
        for pattern in subjective_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return 'subjective'
        
        # 默认为单选题
        return 'single'
    
    def _calculate_confidence(self, text: str, question_type: str) -> float:
        """计算置信度 - 改进版"""
        base_confidence = 0.6
        
        # 根据文本长度调整
        text_length = len(text.strip())
        if text_length < 15:
            base_confidence -= 0.4
        elif text_length < 30:
            base_confidence -= 0.2
        elif text_length > 100:
            base_confidence += 0.2
        elif text_length > 200:
            base_confidence += 0.3
        
        # 根据题型特征调整
        if question_type == 'single':
            # 检查选择项完整性
            choices = re.findall(r'[A-D][.\uff0e、]', text)
            if len(choices) >= 3:
                base_confidence += 0.3
            elif len(choices) >= 2:
                base_confidence += 0.1
                
        elif question_type == 'multiple':
            if '多选' in text or '哪些' in text:
                base_confidence += 0.2
                
        elif question_type == 'fill':
            fill_marks = len(re.findall(r'_{3,}|……|　+', text))
            if fill_marks > 0:
                base_confidence += 0.3
                
        elif question_type == 'judge':
            if re.search(r'\(　*\)|\(\s*\)', text):
                base_confidence += 0.2
                
        elif question_type == 'subjective':
            subjective_words = ['论述', '分析', '简答', '说明', '计算', '证明']
            if any(word in text for word in subjective_words):
                base_confidence += 0.2
        
        # 检查是否包含问号
        if re.search(r'[?？]', text):
            base_confidence += 0.1
        
        # 检查是否包含关键问题词
        question_words = ['什么', '哪个', '哪些', '下列', '以下', '怎么', '为什么', '如何']
        if any(word in text for word in question_words):
            base_confidence += 0.1
        
        # 惩罚无意义内容
        if re.search(r'^　*$|^\s*$|^答题卡|^请将信息', text):
            base_confidence -= 0.5
        
        return min(1.0, max(0.1, base_confidence))
    
    def _suggest_knowledge_points(self, text: str) -> List[Dict[str, Any]]:
        """建议知识点"""
        knowledge_suggestions = []
        
        # 简单的关键词匹配
        keywords = {
            '方程': {'subject': '数学', 'module': '代数', 'point': '方程'},
            '函数': {'subject': '数学', 'module': '代数', 'point': '函数'},
            '三角形': {'subject': '数学', 'module': '几何', 'point': '三角形'},
            '质数': {'subject': '数学', 'module': '数论', 'point': '质数'},
            '阅读': {'subject': '语文', 'module': '阅读理解', 'point': '现代文'},
            'tense': {'subject': '英语', 'module': '语法', 'point': '时态'},
        }
        
        text_lower = text.lower()
        for keyword, kp_info in keywords.items():
            if keyword in text_lower:
                kp_with_confidence = kp_info.copy()
                kp_with_confidence['confidence'] = 0.8  # type: ignore
                knowledge_suggestions.append(kp_with_confidence)
        
        if not knowledge_suggestions:
            knowledge_suggestions.append({
                'subject': '通用',
                'module': '基础',
                'point': '综合',
                'confidence': 0.5
            })
        
        return knowledge_suggestions
    
    def _process_pdf_fallback(self, filename: str) -> Dict[str, Any]:
        """PDF备用处理"""
        mock_text = f"""
1. 下列哪个数是质数？
A. 4  B. 6  C. 7  D. 9

2. 一元二次方程 x² - 5x + 6 = 0 的解是 ______

3. 判断：所有的矩形都是正方形。（  ）

（注：从 {filename} 模拟提取，因缺少PyPDF2库）
        """.strip()
        
        questions = self._smart_split_text(mock_text)
        return {
            'success': True,
            'file_type': 'pdf',
            'items': questions,
            'message': f'从PDF中提取了 {len(questions)} 道题目（模拟）'
        }
    
    def _process_docx_fallback(self, filename: str) -> Dict[str, Any]:
        """DOCX备用处理"""
        mock_text = f"""
1. 下列哪些是质数？（多选）
A. 2  B. 4  C. 7  D. 9  E. 11

2. 填空：x² + 2x - 3 = 0 的解是 x = ______

3. 论述题：请讨论二次函数的性质。

（注：从 {filename} 模拟提取，因缺少python-docx库）
        """.strip()
        
        questions = self._smart_split_text(mock_text)
        return {
            'success': True,
            'file_type': 'docx',
            'items': questions,
            'message': f'从DOCX中提取了 {len(questions)} 道题目（模拟）'
        }